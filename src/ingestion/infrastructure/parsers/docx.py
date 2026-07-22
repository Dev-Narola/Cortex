import io
import uuid

import docx

from src.ingestion.domain.entities import ParsedDocument
from src.ingestion.infrastructure.parser import DocumentParser


class DocxParser(DocumentParser):
    def parse(self, document_id: uuid.UUID, file_bytes: bytes, mime_type: str) -> ParsedDocument:
        if mime_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raise ValueError(f"Unsupported MIME type for DOCX parser: {mime_type}")

        doc = docx.Document(io.BytesIO(file_bytes))
        
        texts = []
        # Extract paragraphs (including headings)
        for p in doc.paragraphs:
            if p.text.strip():
                # Preserve headings structurally
                style_name = p.style.name if p.style else ""
                if "Heading" in style_name:
                    # Convert docx heading levels to markdown-like headers for the chunker
                    level = style_name.replace("Heading", "").strip()
                    prefix = "#" * int(level) if level.isdigit() else "#"
                    texts.append(f"{prefix} {p.text.strip()}")
                else:
                    texts.append(p.text.strip())
                    
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                # Basic tabular extraction
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)
                    
        full_text = "\n\n".join(texts)
        metadata = {
            "parser": "docx"
        }
        
        return ParsedDocument(document_id=document_id, text=full_text, metadata=metadata)
